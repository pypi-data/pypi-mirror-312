import io

from docx import Document


def merge_word_docs_with_tables(
    main_doc_data: bytes,
    table_doc_data: bytes,
    marker_text: str = "TIDOCS_REPLACE_TABLE",
) -> bytes:
    """
    Merges tables from one Word document into another at specified marker locations.

    Args:
        main_doc_data (bytes): The main document binary data
        table_doc_data (bytes): The document containing tables binary data
        marker_text (str): The text to look for where tables should be inserted

    Returns:
        bytes: The merged document as binary data
    """
    # Load both documents from binary data
    main_doc = Document(io.BytesIO(main_doc_data))
    table_doc = Document(io.BytesIO(table_doc_data))

    # Find all tables in the table document
    tables_to_insert = {}
    current_heading = None

    # Associate tables with their preceding headings
    for element in table_doc.element.body:
        if element.tag.endswith("p"):  # It's a paragraph
            paragraph_text = element.text.strip()
            if paragraph_text:
                # print(paragraph_text)
                current_heading = paragraph_text
        elif element.tag.endswith("tbl"):  # It's a table
            if current_heading:
                tables_to_insert[current_heading] = element

    # Process the main document
    for paragraph in main_doc.paragraphs:
        if marker_text in paragraph.text:
            # Get the table associated with this marker
            if paragraph.text in tables_to_insert:
                # Insert table after the paragraph
                table_element = tables_to_insert[paragraph.text]
                paragraph._p.getparent().replace(paragraph._p, table_element)

    # Save the merged document to bytes
    output = io.BytesIO()
    main_doc.save(output)
    return output.getvalue()


# Usage with your existing code
def merge_documents(doc_data: bytes, table_data: bytes) -> bytes:
    """
    Wrapper function to merge your documents using the existing download objects

    Args:
        doc_data (bytes): Main document data from first Pandoc conversion
        table_data (bytes): Table document data from second Pandoc conversion

    Returns:
        bytes: Merged document data
    """
    try:
        merged_data = merge_word_docs_with_tables(doc_data, table_data)
        return merged_data
    except Exception as e:
        print(f"Error merging documents: {str(e)}")
        raise
